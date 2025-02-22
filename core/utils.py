import re
from PIL import Image
import pytesseract
from docx import Document
from langdetect import detect
from pdf2image import convert_from_path
import os
from django.conf import settings


def convert_pdf_to_images(pdf_path):
    """
    Конвертирует PDF в список изображений (по страницам).
    """
    try:
        images = convert_from_path(pdf_path)
        image_paths = []
        for i, image in enumerate(images):
            image_name = f"page_{i + 1}.png"
            image_path = os.path.join(settings.MEDIA_ROOT, image_name)
            image.save(image_path, "PNG")
            image_paths.append(image_path)
        return image_paths
    except Exception as e:
        print(f"Ошибка конвертации PDF: {e}")
        return []


def split_text_by_language(text):
    """
    Разделяет текст на блоки по языкам.
    """
    lines = text.splitlines()
    result = {"rus": [], "eng": [], "spa": []}

    for line in lines:
        try:
            lang = detect(line)
            if lang in result:
                result[lang].append(line)
            else:
                result["eng"].append(line)  # По умолчанию относим к английскому
        except:
            result["eng"].append(line)  # Если язык не определён, считаем английским

    return result


def extract_contacts(text):
    """
    Извлекает email, ссылки и никнеймы из текста.
    """
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    urls = re.findall(r'(https?://[^\s]+)', text)
    telegram_nicks = re.findall(r'@\w+', text)
    return {
        "emails": emails,
        "urls": urls,
        "telegram_nicks": telegram_nicks,
    }


def extract_text_from_image(image_path, lang='eng'):
    """
    Распознает текст из изображения с поддержкой нескольких языков.
    """
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang=lang)
        return text
    except Exception as e:
        print(f"Ошибка OCR: {e}")
        return ""


def save_text_to_txt(file_path, text):
    """
    Сохраняет текст в файл формата .txt.
    """
    try:
        char_count = len(text)  # Подсчёт количества знаков
        header = f"Количество знаков в тексте: {char_count}\n\n"
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(header + text)
    except Exception as e:
        print(f"Ошибка при сохранении в .txt: {e}")


def save_text_to_docx(file_path, text):
    """
    Сохраняет текст в файл формата .docx.
    """
    try:
        # Подсчёт количества знаков
        char_count = len(text)
        header = f"Количество знаков в тексте: {char_count}"
        # Создание документа
        doc = Document()
        doc.add_paragraph(header)  # Добавляем сообщение о количестве знаков
        doc.add_paragraph("")  # Пустая строка для разделения
        doc.add_paragraph(text)  # Добавляем основной текст
        doc.save(file_path)
    except Exception as e:
        print(f"Ошибка при сохранении в .docx: {e}")
